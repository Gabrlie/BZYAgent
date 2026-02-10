"""
授课计划参数解析与规范化工具
"""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document


def extract_text_from_docx_bytes(content: bytes) -> str:
    """从 docx 字节中抽取文本（包含表格）。"""
    doc = Document(BytesIO(content))
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_docx(path: Path) -> str:
    """从 docx 文件中抽取文本。"""
    return extract_text_from_docx_bytes(path.read_bytes())


def extract_text_from_plain_bytes(content: bytes) -> str:
    """从纯文本/markdown 中抽取文本。"""
    return content.decode("utf-8", errors="ignore")


def parse_plan_params_json(raw: Optional[str]) -> Optional[Dict[str, Any]]:
    if not raw:
        return None
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else None
    except Exception:
        return None


def _safe_int(value: Any) -> Optional[int]:
    try:
        if value is None:
            return None
        if isinstance(value, bool):
            return None
        return int(value)
    except Exception:
        return None


def _normalize_tasks(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        parts = [str(item).strip() for item in value if str(item).strip()]
        return "\n".join(parts)
    return str(value).strip()


def normalize_schedule_item(item: Dict[str, Any]) -> Dict[str, Any]:
    week = _safe_int(item.get("week"))
    order = _safe_int(item.get("order") or item.get("sequence") or item.get("lesson_number"))
    hour = _safe_int(item.get("hour") or item.get("hours") or item.get("学时"))
    title = str(item.get("title") or item.get("project_name") or item.get("project") or "").strip()
    tasks = _normalize_tasks(item.get("tasks") or item.get("task") or item.get("content"))

    return {
        "week": week,
        "order": order,
        "hour": hour,
        "title": title,
        "tasks": tasks,
    }


def infer_hour_per_class(schedule: List[Dict[str, Any]], fallback: Optional[int] = None) -> Optional[int]:
    hours = [item.get("hour") for item in schedule if isinstance(item.get("hour"), int) and item.get("hour") > 0]
    if hours:
        # 取众数
        return max(set(hours), key=hours.count)
    return fallback if (isinstance(fallback, int) and fallback > 0) else None


def build_plan_params_from_schedule(
    schedule: List[Dict[str, Any]],
    hour_per_class: Optional[int] = None,
    meta: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    normalized = [normalize_schedule_item(item) for item in schedule if isinstance(item, dict)]
    normalized = [item for item in normalized if item.get("order") is not None]
    normalized.sort(key=lambda item: item.get("order") or 0)

    inferred_hour = infer_hour_per_class(normalized, hour_per_class)
    if inferred_hour:
        for item in normalized:
            if not item.get("hour"):
                item["hour"] = inferred_hour

    params: Dict[str, Any] = {"schedule": normalized}
    if inferred_hour:
        params["hour_per_class"] = inferred_hour
    if meta:
        params.update(meta)
    return params


def build_plan_params_from_content(content: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    if not isinstance(content, dict):
        return None
    schedule = content.get("schedule")
    if not isinstance(schedule, list):
        return None
    return build_plan_params_from_schedule(schedule)


def get_plan_item(schedule: List[Dict[str, Any]], sequence: int) -> Optional[Dict[str, Any]]:
    for item in schedule:
        if item.get("order") == sequence:
            return item
    return None


def compute_cumulative_hours(
    schedule: List[Dict[str, Any]],
    sequence: int,
    default_hour: Optional[int] = None,
) -> int:
    total = 0
    for item in schedule:
        order = item.get("order")
        if isinstance(order, int) and order <= sequence:
            hour = item.get("hour")
            if isinstance(hour, int) and hour > 0:
                total += hour
            elif isinstance(default_hour, int) and default_hour > 0:
                total += default_hour
    return total
